"""
Document Extractors Module

This module provides functionality to extract text from various document formats
including PDF, DOCX, HTML, and TXT files.
"""

import os
import logging
from typing import Dict, List, Optional, Union
from pathlib import Path
import chardet

# Document processing imports
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document
from bs4 import BeautifulSoup

# For Windows compatibility, we'll use file extensions instead of magic
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentExtractor:
    """Base class for document text extraction."""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.html': self._extract_html,
            '.htm': self._extract_html,
            '.txt': self._extract_txt,
        }
    
    def extract_text(self, file_path: Union[str, Path]) -> Dict[str, Union[str, List[str]]]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing:
                - 'text': Extracted text content
                - 'metadata': File metadata
                - 'sentences': List of sentences (if applicable)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        try:
            extractor_func = self.supported_formats[file_extension]
            text = extractor_func(file_path)
            
            # Clean and normalize text
            cleaned_text = self._clean_text(text)
            
            return {
                'text': cleaned_text,
                'metadata': {
                    'file_name': file_path.name,
                    'file_path': str(file_path),
                    'file_size': file_path.stat().st_size,
                    'file_type': file_extension,
                },
                'sentences': self._split_into_sentences(cleaned_text)
            }
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        try:
            text = pdf_extract_text(str(file_path))
            return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            raise
    
    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(str(file_path))
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            return '\n'.join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            raise
    
    def _extract_html(self, file_path: Path) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            soup = BeautifulSoup(content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            return text
        except Exception as e:
            logger.error(f"Error extracting HTML text: {str(e)}")
            raise
    
    def _extract_txt(self, file_path: Path) -> str:
        """Extract text from TXT file."""
        try:
            # Detect encoding
            with open(file_path, 'rb') as file:
                raw_data = file.read()
                encoding_result = chardet.detect(raw_data)
                encoding = encoding_result.get('encoding', 'utf-8')
            
            # Read file with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                text = file.read()
            
            return text
        except Exception as e:
            logger.error(f"Error extracting TXT text: {str(e)}")
            raise
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove special characters but keep punctuation
        import re
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\"\']+', ' ', text)
        
        # Remove multiple spaces
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using simple heuristics."""
        if not text:
            return []
        
        import re
        
        # Simple sentence splitting based on punctuation
        sentences = re.split(r'[.!?]+', text)
        
        # Clean and filter sentences
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:  # Filter out very short fragments
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences
    
    def batch_extract(self, file_paths: List[Union[str, Path]]) -> Dict[str, Dict]:
        """
        Extract text from multiple documents.
        
        Args:
            file_paths: List of file paths to process
            
        Returns:
            Dictionary mapping file paths to extraction results
        """
        results = {}
        
        for file_path in file_paths:
            try:
                results[str(file_path)] = self.extract_text(file_path)
                logger.info(f"Successfully extracted text from: {file_path}")
            except Exception as e:
                logger.error(f"Failed to extract text from {file_path}: {str(e)}")
                results[str(file_path)] = {
                    'text': '',
                    'metadata': {'error': str(e)},
                    'sentences': []
                }
        
        return results


def get_file_type(file_path: Union[str, Path]) -> str:
    """Get the file type using python-magic or fallback to file extension."""
    if HAS_MAGIC:
        try:
            file_type = magic.from_file(str(file_path), mime=True)
            return file_type
        except Exception:
            pass
    
    # Fallback to file extension mapping
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    extension_to_mime = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.html': 'text/html',
        '.htm': 'text/html',
        '.txt': 'text/plain'
    }
    
    return extension_to_mime.get(extension, extension)


# Example usage
if __name__ == "__main__":
    extractor = DocumentExtractor()
    
    # Test with a sample file (if exists)
    sample_file = "sample_docs/sample.txt"
    if os.path.exists(sample_file):
        result = extractor.extract_text(sample_file)
        print(f"Extracted text: {result['text'][:200]}...")
        print(f"Number of sentences: {len(result['sentences'])}")
